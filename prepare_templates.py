"""
Prepara os templates .docx substituindo os campos placeholder
pelos tokens {campo} que o docxtemplater consegue processar.
"""
import zipfile, shutil, re, os

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), 'templates')
PROCESSED_DIR = os.path.join(TEMPLATES_DIR, 'processed')
os.makedirs(PROCESSED_DIR, exist_ok=True)

# ──────────────────────────────────────────────────────────
# Gestor & Líder  (XXXX uppercase, texto com destaques amarelos)
# ──────────────────────────────────────────────────────────
GESTOR_REPLACEMENTS = [
    # Frases completas primeiro (mais específicas)
    ('XXXXXXXXXXXX XXXXXXX XXXXXXX XXXXX', '{razao_social}'),
    ('XXXXXX XXXXX XXXX XXXXXXX XXX',      '{representante}'),
    ('XXXXXX@GMAIL.COM',                   '{email}'),
    ('XXXXXXXXXXX',                        '{cpf}'),
    ('XXXXXXXXX XXXXXX XX XXXX ',          '{endereco} '),
    ('XXXXXXXX',                           '{cnpj}'),
    ('XXXXXXX',                            '{cep}'),
    ('DDD XXXXXX',                         '{telefone}'),
    ('presencial/ home office',            '{modalidade}'),
    # Salário e data de pagamento como frases completas para evitar conflito com "X" solto
    ('R$ XXXXX  (POR EXTENSO)',            'R$ {salario} ({salario_extenso})'),
    ('todo dia X (EXTENSO)',               'todo dia {dia_pagamento} ({dia_pagamento_extenso})'),
    ('Curitiba, X de XXXX de 2026',        'Curitiba, {dia_contrato} de {mes_contrato} de {ano_contrato}'),
    ('CONTRATADA ',                        '{razao_social} '),
]

LIDER_REPLACEMENTS = GESTOR_REPLACEMENTS  # mesma estrutura

# ──────────────────────────────────────────────────────────
# Vendedor & SDR  (xx lowercase, texto fragmentado em runs)
# Usamos contexto para diferenciar campos similares
# ──────────────────────────────────────────────────────────
VENDEDOR_REPLACEMENTS = [
    ('RAZÃO SOCIAL',                            '{razao_social}'),
    ('NOME DO',                                 '{nome_do}'),   # manter split: NOME DO / REPRESENTANTE
    ('REPRESENTANTE',                           '{representante}'),
    # CNPJ: fragmentado em "xx.xxx.xxx" + "/0001-" + "xx"
    # Estratégia: trocar cada fragmento com contexto
    ('xx.xxx.xxx',                              '{cnpj_parte1}'),
    ('/0001-',                                  '/0001-'),       # fixo
    # o "xx" final do CNPJ é tratado junto com o contexto abaixo
    ('Rua',                                     'Rua'),          # fixo, não trocar
    ('Av',                                      'Av'),           # fixo
    ('xxxxx',                                   '{logradouro}'),
    ('nº ',                                     'nº '),          # fixo
    ('bairro ',                                 'bairro '),      # fixo
    ('xxxx',                                    '{cidade_bairro}'),  # bairro ou cidade
    ('xx',                                      '{xx_generico}'),
    ('CEP',                                     'CEP'),
    ('xxx',                                     '{xxx_generico}'),
    ('xxxxxx',                                  '{cpf}'),
    ('eletrônico ',                             'eletrônico '),
    ('xxxxx@gmail.com',                         '{email}'),
    ('WhatsApp ',                               'WhatsApp '),
    ('xxxx-xxxx',                               '{whatsapp_num}'),
    ('XX de',                                   '{dia_contrato} de'),
    ('XXXXX de 2026',                           '{mes_contrato} de 2026'),
]


def remove_yellow_and_fix_bold(xml: str) -> str:
    """
    1. Remove highlight amarelo de todos os runs.
    2. No parágrafo de identificação CONTRATADO/CONTRATADA, o merge coloca
       tudo num único run bold. Divide em: 'CONTRATADO: ' (bold) + resto (normal).
    """
    # ── 1. Remover amarelo ──────────────────────────────────────────────────
    def fix_rpr(m):
        rpr = m.group(0)
        if 'yellow' in rpr:
            rpr = re.sub(r'<w:highlight[^/]*/>', '', rpr)
        # Remove cor vermelha (FF0000) de todos os runs
        rpr = re.sub(r'<w:color w:val="FF0000"\s*/>', '', rpr)
        return rpr
    xml = re.sub(r'<w:rPr>.*?</w:rPr>', fix_rpr, xml, flags=re.DOTALL)

    # ── 2. Dividir run bold do parágrafo de identificação ──────────────────
    # Padrão: run bold com texto "CONTRATADO: ..." ou "CONTRATADA: ..."
    LABELS = ('CONTRATADO: ', 'CONTRATADA: ')

    def split_run(m):
        run = m.group(0)
        # Só age em runs bold
        if not (re.search(r'<w:b\b', run)):
            return run
        # Extrair texto do run
        t_match = re.search(r'(<w:t[^>]*>)(.*?)(</w:t>)', run, re.DOTALL)
        if not t_match:
            return run
        text = t_match.group(2)
        label = next((l for l in LABELS if text.startswith(l)), None)
        if not label or '{razao_social}' not in text:
            return run

        # Extrair rPr do run original
        rpr_match = re.search(r'(<w:rPr>.*?</w:rPr>)', run, re.DOTALL)
        rpr_bold = rpr_match.group(1) if rpr_match else ''
        # rPr sem bold para o restante
        rpr_normal = re.sub(r'<w:b\b[^/]*/>', '', rpr_bold)
        rpr_normal = re.sub(r'<w:bCs\b[^/]*/>', '', rpr_normal)

        rest = text[len(label):]
        # Run 1: label em bold
        run1 = re.sub(r'<w:t[^>]*>.*?</w:t>', f'<w:t xml:space="preserve">{label}</w:t>', run, flags=re.DOTALL)
        # Run 2: restante sem bold
        run2 = re.sub(r'<w:rPr>.*?</w:rPr>', rpr_normal, run, flags=re.DOTALL) if rpr_bold else run
        run2 = re.sub(r'<w:t[^>]*>.*?</w:t>', f'<w:t xml:space="preserve">{rest}</w:t>', run2, flags=re.DOTALL)
        return run1 + run2

    xml = re.sub(r'<w:r[ >].*?</w:r>', split_run, xml, flags=re.DOTALL)
    return xml


def replace_text_in_xml(xml: str, old: str, new: str) -> str:
    """Substitui texto diretamente no conteúdo de tags <w:t>."""
    # Substitui SOMENTE dentro de tags <w:t>...</w:t>
    def replacer(m):
        return m.group(0).replace(old, new)
    return re.sub(r'<w:t[^>]*>.*?</w:t>', replacer, xml, flags=re.DOTALL)


def process_docx(input_path: str, output_path: str, replacements: list):
    shutil.copy2(input_path, output_path)

    # Ler todos os arquivos do ZIP
    with zipfile.ZipFile(input_path, 'r') as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    xml = files['word/document.xml'].decode('utf-8')

    for old, new in replacements:
        xml = xml.replace(old, new)

    files['word/document.xml'] = xml.encode('utf-8')

    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)

    print(f'  ✓ {os.path.basename(output_path)}')


# ──────────────────────────────────────────────────────────
# Processamento especial para Vendedor/SDR (runs fragmentados)
# Merge de runs dentro do mesmo parágrafo antes de substituir
# ──────────────────────────────────────────────────────────
def replace_in_para(para, field_map: dict):
    """Merge todos os runs (incluindo dentro de hyperlinks) e faz substituições."""
    from lxml import etree
    WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def get_all_runs(element):
        """Retorna todos os <w:r> dentro do elemento, incluindo hyperlinks."""
        return element.findall(f'.//{{{WNS}}}r')

    all_runs = get_all_runs(para._element)
    if not all_runs:
        return

    full = ''.join(
        (r.find(f'{{{WNS}}}t').text or '') if r.find(f'{{{WNS}}}t') is not None else ''
        for r in all_runs
    )

    new_full = full
    for old, new in field_map.items():
        new_full = new_full.replace(old, new)

    if new_full != full:
        # Coloca tudo no primeiro run, zera os demais
        t0 = all_runs[0].find(f'{{{WNS}}}t')
        if t0 is not None:
            t0.text = new_full
            t0.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        for r in all_runs[1:]:
            t = r.find(f'{{{WNS}}}t')
            if t is not None:
                t.text = ''


def merge_runs_and_replace(input_path: str, output_path: str, field_map: dict):
    """
    Para templates com runs fragmentados (incluindo dentro de hyperlinks):
    - Para cada parágrafo, concatena o texto de todos os runs
    - Faz substituições no texto completo
    - Reescreve apenas o primeiro run com o texto resultante, zerando os demais
    """
    try:
        from docx import Document
    except ImportError:
        print('  ⚠️  python-docx não instalado, usando substituição direta')
        process_docx(input_path, output_path, list(field_map.items()))
        return

    doc = Document(input_path)

    for para in doc.paragraphs:
        replace_in_para(para, field_map)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para, field_map)

    doc.save(output_path)

    # Pós-processamento: remover amarelo e corrigir bold via XML
    with zipfile.ZipFile(output_path, 'r') as zin:
        files = {name: zin.read(name) for name in zin.namelist()}
    xml = files['word/document.xml'].decode('utf-8')
    xml = remove_yellow_and_fix_bold(xml)
    files['word/document.xml'] = xml.encode('utf-8')
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)

    print(f'  ✓ {os.path.basename(output_path)} (merge mode)')


# ──────────────────────────────────────────────────────────
# Mapeamentos finais para Vendedor / SDR
# ──────────────────────────────────────────────────────────
VENDEDOR_MAP = {
    # Mais específicos primeiro
    'serviços de Representante de Vendas (Closer)': 'serviços de {nome_cargo}',
    'CONTRATADA: RAZÃO SOCIAL':   'CONTRATADA: {razao_social}',  # corpo do contrato
    'RAZÃO SOCIAL':               '{representante}',              # assinatura (standalone)
    'NOME DO REPRESENTANTE':      '{representante}',
    'xx.xxx.xxx/0001-xx':         '{cnpj}',
    'Rua/Av xxxxx':               'Rua/Av {logradouro}',
    ', nº xx,':                   ', nº {numero},',
    'bairro xxxx,':               'bairro {bairro},',
    'xxxx/xx, CEP:':              '{cidade}/{uf}, CEP:',
    'xx.xxx-xxx':                 '{cep}',
    'CPF sob o nº xxxxxx':        'CPF sob o nº {cpf}',
    'xxxxx@gmail.com':            '{email}',
    '(xx) xxxx-xxxx':             '{whatsapp}',
    'Curitiba, XX de XXXXX de 2026': 'Curitiba, {dia_contrato} de {mes_contrato} de {ano_contrato}',
    '§1°- Os valores apurados podem ser pagos via boleto, PIX, transferência ou depósito bancário, conforme preferência informada pela CONTRATADA e autorizada pela CONTRATANTE.':
        '§1°- Os valores apurados podem ser pagos via PIX, transferência ou depósito bancário, por meio da chave PIX {chave_pix}.',
    'até o dia 20 (vinte) do mês subsequente': 'até o dia {dia_pagamento} ({dia_pagamento_extenso}) do mês subsequente',
}

SDR_MAP = {
    **VENDEDOR_MAP,
    'CHAVE PIX': '{chave_pix}',
    'todo dia 30 (trinta), posteriormente': 'todo dia {dia_pagamento} ({dia_pagamento_extenso}), posteriormente',
}


GESTOR_MAP = {
    'serviços de Líder Técnico':          'serviços de {nome_cargo}',
    'CONTRATADA ':                        '{representante}',
    'XXXXXXXXXXXX XXXXXXX XXXXXXX XXXXX': '{razao_social}',
    'XXXXXX XXXXX XXXX XXXXXXX XXX':      '{representante}',
    'XXXXXX@GMAIL.COM':                   '{email}',
    'XXXXXXXXXXX':                        '{cpf}',
    'XXXXXXXXX XXXXXX XX XXXX ':          '{endereco} ',
    'XXXXXXXX':                           '{cnpj}',
    'XXXXXXX':                            '{cep}',
    'DDD XXXXXX':                         '{telefone}',
    'presencial/ home office':            '{modalidade}',
    # Salário: XXXXX + POR EXTENSO em runs separados — substituir individualmente
    'XXXXX':                              '{salario}',
    'POR EXTENSO':                        '{salario_extenso}',
    'todo dia X (EXTENSO)':               'todo dia {dia_pagamento} ({dia_pagamento_extenso})',
    '§1°- Os valores serão pagos via PIX, transferência ou depósito bancário.':
        '§1°- Os valores serão pagos via PIX, transferência ou depósito bancário, por meio da chave vinculada ao {chave_pix}.',
    'Curitiba, X de XXXX de 2026':        'Curitiba, {dia_contrato} de {mes_contrato} de {ano_contrato}',
}

LIDER_MAP = dict(GESTOR_MAP)


if __name__ == '__main__':
    print('Preparando templates...')

    merge_runs_and_replace(
        os.path.join(TEMPLATES_DIR, 'gestor.docx'),
        os.path.join(PROCESSED_DIR, 'gestor.docx'),
        GESTOR_MAP
    )
    merge_runs_and_replace(
        os.path.join(TEMPLATES_DIR, 'lider.docx'),
        os.path.join(PROCESSED_DIR, 'lider.docx'),
        LIDER_MAP
    )
    merge_runs_and_replace(
        os.path.join(TEMPLATES_DIR, 'vendedor.docx'),
        os.path.join(PROCESSED_DIR, 'vendedor.docx'),
        VENDEDOR_MAP
    )
    merge_runs_and_replace(
        os.path.join(TEMPLATES_DIR, 'sdr.docx'),
        os.path.join(PROCESSED_DIR, 'sdr.docx'),
        SDR_MAP
    )

    print('\nTemplates prontos em templates/processed/')
